#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import norm, shapiro, normaltest, kstest
import os
import base64

from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm

# 1) Configuração da página (deve ser a primeira chamada de Streamlit)
st.set_page_config(page_title="Consumo Referencial", layout="centered")

# 2) Persistência do DataFrame e chave do uploader no session_state
if "df_consumo" not in st.session_state:
    st.session_state.df_consumo = None
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

# 3) Submenu "Abastecimento de Água" com as quatro opções
st.sidebar.title("Demanda Hídrica:")
aba = st.sidebar.selectbox("Consumo e Vazão", [
    "🧮 Cálculo",
    "📊 Gerar Histograma",
    "ℹ️ Sobre esse App",
    "📘 Sobre o Modelo Estatístico"
])

# Função para formatação numérica: ponto para milhar e vírgula para decimal
def format_num(value, decimals=2):
    fmt = f",.{decimals}f" if decimals > 0 else f",.0f"
    return f"{value:{fmt}}".replace(",", "X").replace(".", ",").replace("X", ".")

# 4) Aba "Cálculo do Consumo e Vazão"
if aba == "🧮 Cálculo":
    st.title("Cálculo do Consumo Referencial")

    st.header("Dados do Projeto")
    # Limite de 140 caracteres para os campos de texto
    nome_projeto = st.text_input("Nome do Projeto", value="Projeto 1", max_chars=140)
    tecnico_operador = st.text_input("Técnico Operador", value="", max_chars=140)
    tipo_medicao = st.selectbox("Tipo de Medição", [
        "Micromedição - Hidrômetros",
        "Macromedição - Sensores de Vazão"
    ])

    st.header("1. Dados de Consumo Mensal")
    # Se o CSV já estiver carregado, exibe mensagem e permite carregar outro
    if st.session_state.df_consumo is not None:
        st.info("Arquivo CSV já carregado.")
        if st.button("Carregar outro arquivo CSV"):
            st.session_state.df_consumo = None
            st.session_state.uploader_key += 1  # Reinicializa o uploader
            pass
    else:
        uploaded_file = st.file_uploader(
            "Faça o upload de um arquivo CSV (2 colunas: Mês, Consumo (m³))",
            type="csv",
            key=st.session_state.uploader_key
        )
        if uploaded_file is not None:
            try:
                df = pd.read_csv(uploaded_file)
                if df.shape[1] >= 2:
                    df.columns = ['Mês', 'Consumo (m³)']
                    st.session_state.df_consumo = df
                    st.success("Arquivo carregado com sucesso!")
                else:
                    st.error("O arquivo precisa ter pelo menos duas colunas.")
            except Exception as e:
                st.error(f"Erro ao ler o CSV: {e}")

    # Se o CSV está carregado, prossegue com o cálculo
    if st.session_state.df_consumo is not None:
        df = st.session_state.df_consumo
        st.dataframe(df)

        st.header("2. Parâmetros do Projeto")
        modelo = st.selectbox("Modelo Estatístico", ["KDE", "Distribuição Normal"])
        percentil = st.slider("Percentil de Projeto (%)", 50, 99, 95)
        dias_mes = st.number_input("Número de dias do mês", min_value=1, max_value=31, value=30)

        # Novo campo: Número de horas diárias de operação (1 <= t <= 24)
        horas_operacao = st.number_input("Número de horas diárias de operação", min_value=1, max_value=24, value=24, step=1)

        tempo_dia = 86400  # Valor fixo (segundos em um dia)
        k1 = st.number_input("Coeficiente de máx. diária (K1)", min_value=1.0, value=1.4)
        k2 = st.number_input("Coeficiente de máx. horária (K2)", min_value=1.0, value=2.0)

        consumo = df['Consumo (m³)'].values

        # Cálculo do consumo referencial
        if modelo == "KDE":
            fig_temp, ax_temp = plt.subplots()
            kde = sns.kdeplot(consumo, bw_adjust=1, ax=ax_temp)
            kde_y = kde.get_lines()[0].get_ydata()
            kde_x = kde.get_lines()[0].get_xdata()
            plt.close(fig_temp)
            cdf_kde = np.cumsum(kde_y)
            cdf_kde = cdf_kde / cdf_kde[-1]
            consumo_ref = kde_x[np.searchsorted(cdf_kde, percentil / 100)]
        else:
            consumo_ref = np.percentile(consumo, percentil)

        # Fator de ajuste (r = 24 / t)
        r = 24 / horas_operacao

        # Cálculo das vazões (com fator de ajuste)
        q_med_base = (consumo_ref / dias_mes) / tempo_dia * 1000
        q_med = q_med_base * r
        q_max_dia = q_med * k1
        q_max_hora = q_med * k2
        q_max_real = q_med * k1 * k2

        desvio_padrao = np.std(consumo)
        media = np.mean(consumo)

        st.header("3. Resultados")
        # Exibição em 3 colunas e 2 linhas
        col1, col2, col3 = st.columns(3)
        # Primeira linha
        col1.metric("Consumo Referencial (m³)", format_num(consumo_ref, 0))
        col2.metric("Desvio Padrão (m³)", format_num(desvio_padrao, 2))
        col3.metric("Vazão Média (L/s)", format_num(q_med, 2))
        # Segunda linha
        col1.metric("Vazão Máx. Diária (L/s)", format_num(q_max_dia, 2))
        col2.metric("Vazão Máx. Horária (L/s)", format_num(q_max_hora, 2))
        col3.metric("Vazão Máx. Dia+Hora (L/s)", format_num(q_max_real, 2))

        st.header("4. Testes de Normalidade")
        stat_sw, p_sw = shapiro(consumo)
        stat_dp, p_dp = normaltest(consumo)
        stat_ks, p_ks = kstest(consumo, 'norm', args=(media, desvio_padrao))

        def interpreta(p):
            return "✔️ Aceita a hipótese de normalidade." if p > 0.05 else "❌ Rejeita a hipótese de normalidade."

        txt_sw = f"Shapiro-Wilk: Estatística = {format_num(stat_sw, 3)}; p-valor = {format_num(p_sw, 3)} — {interpreta(p_sw)}"
        txt_dp = f"D'Agostino e Pearson: Estatística = {format_num(stat_dp, 3)}; p-valor = {format_num(p_dp, 3)} — {interpreta(p_dp)}"
        txt_ks = f"Kolmogorov-Smirnov (KS): Estatística = {format_num(stat_ks, 3)}; p-valor = {format_num(p_ks, 3)} — {interpreta(p_ks)}"

        st.write(f"**{txt_sw}**")
        st.write(f"**{txt_dp}**")
        st.write(f"**{txt_ks}**")

        # Novo campo para escolha da apresentação do histograma
        tipo_hist = st.selectbox("Tipo de apresentação do histograma:", ["Frequência Absoluta", "Densidade de Probabilidade"])
        if tipo_hist == "Densidade de Probabilidade":
            stat_param = "density"
        else:
            stat_param = "count"

        st.header("5. Gráfico de Distribuição")
        fig1, ax1 = plt.subplots(figsize=(10, 5))
        sns.histplot(consumo, kde=True, stat=stat_param, color="skyblue", edgecolor="black", bins=12, ax=ax1)
        x_vals = np.linspace(min(consumo), max(consumo), 1000)
        # Se o histograma for em frequência absoluta, escalamos a curva normal
        if stat_param == "count":
            bin_width = (max(consumo) - min(consumo)) / 12
            n = len(consumo)
            normal_curve = norm.pdf(x_vals, loc=media, scale=desvio_padrao) * n * bin_width
        else:
            normal_curve = norm.pdf(x_vals, loc=media, scale=desvio_padrao)
        ax1.plot(x_vals, normal_curve, color='red', linestyle='--', label='Distribuição Normal')
        ax1.axvline(consumo_ref, color='black', linestyle=':', label=f'{percentil}% ≈ {format_num(consumo_ref, 0)} m³')
        ax1.set_xlabel("Consumo mensal (m³)")
        ax1.set_ylabel("Frequência" if stat_param=="count" else "Densidade estimada")
        ax1.set_title("Distribuição do Consumo com KDE e Normal")
        ax1.legend()
        st.pyplot(fig1)

        st.header("6. Funções de Distribuição Acumulada")
        fig_temp2, ax_temp2 = plt.subplots()
        kde_plot = sns.kdeplot(consumo, bw_adjust=1, ax=ax_temp2)
        kde_y2 = kde_plot.get_lines()[0].get_ydata()
        kde_x2 = kde_plot.get_lines()[0].get_xdata()
        plt.close(fig_temp2)
        cdf_kde2 = np.cumsum(kde_y2)
        cdf_kde2 = cdf_kde2 / cdf_kde2[-1]
        cdf_norm = norm.cdf(kde_x2, loc=media, scale=desvio_padrao)

        fig2, ax2 = plt.subplots(figsize=(8, 5))
        ax2.plot(kde_x2, cdf_kde2, label='CDF da KDE', color='blue')
        ax2.plot(kde_x2, cdf_norm, label='CDF da Normal', color='red', linestyle='--')
        ax2.set_title("Funções de Distribuição Acumulada (CDF) KDE vs Distribuição Normal")
        ax2.set_xlabel("Consumo mensal de água (m³)")
        ax2.set_ylabel("Probabilidade acumulada")
        ax2.legend()
        ax2.grid(True)
        st.pyplot(fig2)

        st.header("Relatório em Word")
        if st.button("Gerar Relatório Word"):
            doc = Document()
            # Ajuste das margens: superior e inferior = 2 cm; esquerda e direita = 2,5 cm
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            
            doc.add_heading("Relatório de Consumo Referencial", 0)
            doc.add_heading("Dados do Projeto", level=1)
            doc.add_paragraph(f"Nome do Projeto: {nome_projeto}")
            doc.add_paragraph(f"Técnico Operador: {tecnico_operador}")
            doc.add_paragraph(f"Tipo de Medição: {tipo_medicao}")
            doc.add_heading("Parâmetros de Entrada", level=1)
            doc.add_paragraph(f"Modelo Estatístico: {modelo}")
            doc.add_paragraph(f"Percentil de Projeto: {percentil}%")
            doc.add_paragraph(f"Número de dias do mês: {dias_mes}")
            doc.add_paragraph(f"Número de horas diárias de operação: {horas_operacao}")
            doc.add_paragraph(f"Tempo diário (s): {tempo_dia}")
            doc.add_paragraph(f"K1 (máx. diária): {format_num(k1, 2)}")
            doc.add_paragraph(f"K2 (máx. horária): {format_num(k2, 2)}")
            doc.add_heading("Resultados", level=1)
            doc.add_paragraph(f"Consumo Referencial (m³): {format_num(consumo_ref, 0)}")
            doc.add_paragraph(f"Desvio Padrão (m³): {format_num(desvio_padrao, 2)}")
            doc.add_paragraph(f"Vazão Média (L/s): {format_num(q_med, 2)}")
            doc.add_paragraph(f"Vazão Máx. Diária (L/s): {format_num(q_max_dia, 2)}")
            doc.add_paragraph(f"Vazão Máx. Horária (L/s): {format_num(q_max_hora, 2)}")
            doc.add_paragraph(f"Vazão Máx. Dia+Hora (L/s): {format_num(q_max_real, 2)}")
            doc.add_heading("Testes de Normalidade", level=1)
            doc.add_paragraph(txt_sw)
            doc.add_paragraph(txt_dp)
            doc.add_paragraph(txt_ks)
            img_buffer1 = BytesIO()
            fig1.savefig(img_buffer1, format="png", dpi=150)
            img_buffer1.seek(0)
            doc.add_heading("Gráfico de Distribuição", level=1)
            doc.add_picture(img_buffer1, width=Inches(6))
            img_buffer2 = BytesIO()
            fig2.savefig(img_buffer2, format="png", dpi=150)
            img_buffer2.seek(0)
            doc.add_heading("Funções de Distribuição Acumulada", level=1)
            doc.add_picture(img_buffer2, width=Inches(6))
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            st.download_button(
                label="Baixar Relatório Word",
                data=doc_buffer.getvalue(),
                file_name="Relatorio_Consumo.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# 5) Aba "Gerar Histograma"
elif aba == "📊 Gerar Histograma":
    st.title("Gerar Tabela de Consumo Mensal")
    st.markdown("Informe os dados do projeto para gerar uma planilha de consumo mensal de água tratada.")
    ano_inicial = st.number_input("Ano Inicial", min_value=2000, max_value=2100, value=2020, step=1)
    ano_final = st.number_input("Ano Final", min_value=2000, max_value=2100, value=2025, step=1)
    populacao = st.number_input("População atendida", min_value=1000, value=80000, step=1000)
    if st.button("Criar Planilha"):
        if ano_final < ano_inicial:
            st.error("O Ano Final deve ser maior ou igual ao Ano Inicial.")
        else:
            anos = list(range(int(ano_inicial), int(ano_final) + 1))
            meses = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
                     "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
            fator = populacao / 50000.0
            media_baseline = 300000 * fator
            std_baseline = 50000 * fator
            registros = []
            for ano in anos:
                for mes in meses:
                    consumo_val = int(np.random.normal(media_baseline, std_baseline))
                    mes_ano = f"{mes}/{ano}"
                    registros.append({"Mês": mes_ano, "Consumo (m³)": consumo_val})
            df_gerado = pd.DataFrame(registros)
            st.dataframe(df_gerado)
            csv_gerado = df_gerado.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="Baixar Planilha CSV",
                data=csv_gerado,
                file_name="Consumo_Mensal_Agua_Tratada.csv",
                mime="text/csv"
            )

# 6) Aba "ℹ️ Sobre esse App"
elif aba == "ℹ️ Sobre esse App":
    st.title("Sobre esse App")
    # HTML com estilo unificado (fonte Arial, tamanho 16, espaçamento 1.5)
    # Adicionado li { margin-bottom: 2em; } para espaçamento duplo entre itens numerados
    html_content = """
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
      <meta charset="UTF-8" />
      <title>Sobre esse App</title>
      <style>
        body, h1, h2, h3, p, ol, ul, li {
          font-family: "Arial", sans-serif;
          font-size: 16px;
          line-height: 1.5;
        }
        li {
          margin-bottom: 2em; /* Espaçamento duplo entre itens numerados */
        }
        .page { 
          display: none; 
          margin: 20px; 
        }
        .page.active { 
          display: block; 
        }
        .nav-buttons { 
          margin-top: 20px; 
        }
        button { 
          margin: 5px; 
          padding: 8px 16px; 
          cursor: pointer; 
        }
        code {
          background-color: #f5f5f5; 
          padding: 2px 4px; 
          font-size: 90%; 
          border-radius: 4px; 
          font-family: Consolas, monospace;
        }
      </style>
      <!-- MathJax para renderizar LaTeX -->
      <script>
        window.MathJax = {
          tex: { inlineMath: [['$', '$'], ['\\(', '\\)']] }
        };
      </script>
      <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-chtml.js"></script>
    </head>
    <body>
      <!-- Página 1 -->
      <div class="page active" id="page1">
        <h1>Sobre este Aplicativo</h1>
        <p>
          Este aplicativo foi desenvolvido em <code>Python</code> utilizando a biblioteca <code>Streamlit</code> 
          para análise estatística de consumo de água. Ele permite o carregamento de arquivos CSV com dados de consumo, 
          realiza cálculos estatísticos, gera gráficos e produz relatórios. Além do consumo mensal referencial, 
          o aplicativo também calcula as vazões médias em função dos dias de operação do sistema, possibilitando 
          avaliar a variação dessas vazões como uma estimativa preliminar para análise do projetista.
        </p>
        <p>
          Para maior flexibilidade, foi incluído um fator de ajuste baseado no número de horas diárias de operação 
          do sistema, permitindo ajustar as equações de vazão. Dessa forma, é possível simular diferentes cenários 
          de operação entre 1 hora e 24 horas diárias.
        </p>
        <div class="nav-buttons">
          <button onclick="showPage(2)">Próxima &raquo;</button>
        </div>
      </div>

      <!-- Página 2 -->
      <div class="page" id="page2">
        <h2>Estrutura do Código</h2>
        <ol>
          <li>
            <strong>Importações e Configurações:</strong> Importa bibliotecas como 
            <code>pandas</code>, <code>numpy</code>, <code>matplotlib</code>, <code>seaborn</code> 
            e faz a chamada <code>st.set_page_config</code> logo no início, sendo a primeira instrução de Streamlit.
          </li>
          <li>
            <strong>Session State:</strong> Utiliza <code>st.session_state</code> para manter dados entre interações.
          </li>
          <li>
            <strong>Menu de Navegação:</strong> Define as funcionalidades do app, como 
            <em>Cálculo do Consumo</em>, <em>Gerar Histograma</em>, <em>Sobre esse App</em> e 
            <em>Sobre o Modelo Estatístico</em>.
          </li>
          <li>
            <strong>Cálculo do Consumo:</strong> Permite o upload do CSV, configura parâmetros, incluindo número 
            de horas diárias de operação, executa cálculos estatísticos e gera gráficos.
          </li>
          <li>
            <strong>Relatório em Word:</strong> Gera um documento com os resultados e gráficos utilizando 
            a biblioteca <code>python-docx</code>.
          </li>
        </ol>
        <div class="nav-buttons">
          <button onclick="showPage(1)">&laquo; Anterior</button>
          <button onclick="showPage(3)">Próxima &raquo;</button>
        </div>
      </div>

      <!-- Página 3 -->
      <div class="page" id="page3">
        <h2>Cálculos e Equações</h2>
        <p>
          Um dos cálculos principais é a determinação do consumo referencial. Para o modelo de distribuição normal, usamos:
          $$ f(x) = \\frac{1}{\\sigma\\sqrt{2\\pi}} \\exp\\Bigl(-\\frac{(x-\\mu)^2}{2\\sigma^2}\\Bigr). $$
        </p>
        <p>
          Os testes de normalidade: Shapiro-Wilk, D'Agostino-Pearson e Kolmogorov-Smirnov, que verificam se os dados 
          seguem uma distribuição normal, aceitando a hipótese quando 
          $$ p\\text{-valor} > 0.05. $$
        </p>
        <p>
          Além disso, o fator de ajuste <em>r</em> é dado por 
          $$ r = \\frac{24}{t}, $$
          onde <em>t</em> é o número de horas diárias de operação entre 1 e 24 horas. 
          Esse fator multiplica as equações de vazão, permitindo avaliar cenários de operação em períodos reduzidos, como 
          por exemplo, apenas 8 horas por dia ou período integral de 24 horas.
        </p>
        <div class="nav-buttons">
          <button onclick="showPage(2)">&laquo; Anterior</button>
          <button onclick="showPage(4)">Próxima &raquo;</button>
        </div>
      </div>

      <!-- Página 4 -->
      <div class="page" id="page4">
        <h2>Equações de Vazão e Variáveis</h2>
        <p>
          Para calcular as vazões, definimos inicialmente:
        </p>
        <ul>
          <li><strong>Consumo Referencial</strong>: valor estatístico que representa o consumo mensal alvo (m³).</li>
          <li><strong>Dias de Operação</strong> (<em>dias_mes</em>): número de dias no mês considerado.</li>
          <li><strong>Horas de Operação</strong> (<em>t</em>): quantas horas por dia o sistema fica operando.</li>
          <li><strong>Fator de Ajuste</strong> (<em>r</em>): $$r = \\frac{24}{t}.$$</li>
          <li><strong>k1</strong>: coeficiente de máxima diária.</li>
          <li><strong>k2</strong>: coeficiente de máxima horária.</li>
        </ul>
        <p>
          A vazão média básica (<em>q_med_base</em>) é calculada por:
        </p>
        <p>
          $$ q_{\\text{med\\_base}} = \\frac{\\text{Consumo Referencial}}{\\text{dias\\_mes}} 
          \\times \\frac{1}{\\text{tempo\\_dia}} \\times 1000. $$
        </p>
        <p>
          Em seguida, aplicamos o fator <em>r</em> para obter a vazão média final:
        </p>
        <p>
          $$ q_{\\text{med}} = q_{\\text{med\\_base}} \\times r. $$
        </p>
        <p>
          As demais vazões são:
        </p>
        <ul>
          <li><em>Vazão Máx. Diária</em>: $$ q_{\\text{max\\_dia}} = q_{\\text{med}} \\times k1. $$</li>
          <li><em>Vazão Máx. Horária</em>: $$ q_{\\text{max\\_hora}} = q_{\\text{med}} \\times k2. $$</li>
          <li><em>Vazão Máx. Dia+Hora</em>: $$ q_{\\text{max\\_real}} = q_{\\text{med}} \\times k1 \\times k2. $$</li>
        </ul>
        <div class="nav-buttons">
          <button onclick="showPage(3)">&laquo; Anterior</button>
        </div>
      </div>

      <script>
        function showPage(pageNumber) {
          document.getElementById("page1").classList.remove("active");
          document.getElementById("page2").classList.remove("active");
          document.getElementById("page3").classList.remove("active");
          document.getElementById("page4").classList.remove("active");
          document.getElementById("page" + pageNumber).classList.add("active");
        }
      </script>
    </body>
    </html>
    """
    st.components.v1.html(html_content, height=800, scrolling=True)

# 7) Aba "📘 Sobre o Modelo Estatístico"
elif aba == "📘 Sobre o Modelo Estatístico":
    st.title("📘 Sobre o Modelo Estatístico")
    pdf_file = "03_Estatistica_2025.pdf"
    if os.path.exists(pdf_file):
        with open(pdf_file, "rb") as f:
            pdf_bytes = f.read()
        st.download_button(
            label="Baixar Relatório PDF",
            data=pdf_bytes,
            file_name="03_Estatistica_2025.pdf",
            mime="application/pdf"
        )
    else:
        st.warning(f"Arquivo PDF '{pdf_file}' não encontrado no diretório atual.")

